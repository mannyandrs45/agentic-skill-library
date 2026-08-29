#!/usr/bin/env python3
"""
Resume DOCX Generator
Generates a DOCX resume matching the canonical template formatting.

Usage:
    python generate_docx.py <input_json> <output_docx>

Input JSON schema:
{
    "name": "Full Name",
    "contact": "City, ST | (XXX) XXX-XXXX | email@domain.com | linkedin.com/in/handle",
    "education": [
        {
            "institution": "University Name",
            "location": "City, ST",
            "date": "May 2025",
            "lines": [
                {"type": "degree", "text": "Bachelor of Arts | Major: Economics | Minors: Psychology, Data Science"},
                {"type": "honors", "text": "Award Name Here"}
            ]
        }
    ],
    "professional_experience": [
        {
            "company": "Company Name",
            "role": "Job Title",
            "location": "City, ST",
            "dates": "Sep 2025 – Present",
            "bullets": [
                "Achievement bullet 1",
                "Achievement bullet 2"
            ]
        }
    ],
    "leadership_experience": [
        {
            "organization": "Org Name",
            "role": "Role Title",
            "location": "City, ST",
            "dates": "Jan 2023 – Dec 2024",
            "bullets": [
                "Impact bullet 1",
                "Impact bullet 2"
            ]
        }
    ],
    "skills_interests": {
        "technical": "Skill A, Skill B, Skill C",
        "languages": "Spanish (native), Portuguese (professional working proficiency)",
        "interests": "Interest A; Interest B; Interest C"
    }
}
"""

import json
import sys
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.0):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_right_tab(paragraph, tab_position_inches=7.0):
    """Add a right-aligned tab stop to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(int(tab_position_inches * 1440)))
    tabs.append(tab)
    pPr.append(tabs)


def add_run(paragraph, text, bold=False, italic=False, size=10):
    """Add a run to a paragraph with specified formatting."""
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def generate_resume(data, output_path):
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Emu(450850)
        section.bottom_margin = Emu(593725)
        section.left_margin = Emu(456565)
        section.right_margin = Emu(436245)

    # Remove default empty paragraph
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # === HEADER: Name ===
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_para, before=0, after=0)
    add_run(name_para, data["name"], bold=True, size=14)

    # === HEADER: Contact ===
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact_para, before=3.35, after=0)
    add_run(contact_para, data["contact"], size=9)

    # === Blank line after header ===
    blank = doc.add_paragraph()
    set_paragraph_spacing(blank, before=0, after=0)

    # === EDUCATION ===
    edu_heading = doc.add_paragraph()
    set_paragraph_spacing(edu_heading, before=3.35, after=0)
    add_run(edu_heading, "EDUCATION", bold=True, size=12)

    for school in data["education"]:
        # Institution line with right-aligned date
        inst_para = doc.add_paragraph()
        set_paragraph_spacing(inst_para, before=3.35, after=0)
        add_right_tab(inst_para)
        add_run(inst_para, school["institution"] + " ", bold=True, size=10)
        add_run(inst_para, "| " + school["location"], size=10)
        add_run(inst_para, "\t", size=10)
        add_run(inst_para, school["date"], size=10)

        # Sub-lines (degree, honors, scores, etc.)
        for line in school.get("lines", []):
            line_para = doc.add_paragraph()
            set_paragraph_spacing(line_para, before=3.35, after=0)

            if line["type"] == "degree":
                # Parse degree text: "Bachelor of Arts | Major: Economics | Minors: X, Y"
                parts = line["text"].split(" | ")
                for idx, part in enumerate(parts):
                    if idx > 0:
                        add_run(line_para, " | ", size=10)
                    # Check if part has a label like "Major:" or "Minors:"
                    if ":" in part:
                        label, value = part.split(":", 1)
                        add_run(line_para, label + ":", italic=True, size=10)
                        add_run(line_para, value, size=10)
                    else:
                        # Degree name itself
                        add_run(line_para, part, italic=True, size=10)

            elif line["type"] in ("honors", "awards"):
                add_run(line_para, "Honors/Awards: ", italic=True, size=10)
                add_run(line_para, line["text"], size=10)

            elif line["type"] == "scores":
                add_run(line_para, "ACT Scores: ", italic=True, size=10)
                add_run(line_para, line["text"], size=10)

            elif line["type"] == "gpa":
                add_run(line_para, "GPA: ", italic=True, size=10)
                add_run(line_para, line["text"], size=10)

            else:
                # Generic line
                add_run(line_para, line["text"], size=10)

    # === Blank line ===
    blank2 = doc.add_paragraph()
    set_paragraph_spacing(blank2, before=0, after=0)

    # === PROFESSIONAL EXPERIENCE ===
    prof_heading = doc.add_paragraph()
    set_paragraph_spacing(prof_heading, before=3.35, after=0)
    add_run(prof_heading, "PROFESSIONAL EXPERIENCE", bold=True, size=12)

    for job in data["professional_experience"]:
        role_para = doc.add_paragraph()
        set_paragraph_spacing(role_para, before=3.35, after=0)
        add_right_tab(role_para)
        add_run(role_para, job["company"], bold=True, size=10)
        add_run(role_para, ", ", size=10)
        add_run(role_para, job["role"] + " ", italic=True, size=10)
        add_run(role_para, "| " + job["location"], size=10)
        add_run(role_para, "\t", size=10)
        add_run(role_para, job["dates"], size=10)

        for bullet in job.get("bullets", []):
            bp = doc.add_paragraph()
            set_paragraph_spacing(bp, before=3.35, after=0)
            add_run(bp, bullet, size=10)

    # === Blank line ===
    blank3 = doc.add_paragraph()
    set_paragraph_spacing(blank3, before=0, after=0)

    # === LEADERSHIP & OTHER EXPERIENCE ===
    lead_heading = doc.add_paragraph()
    set_paragraph_spacing(lead_heading, before=3.35, after=0)
    add_run(lead_heading, "LEADERSHIP & OTHER EXPERIENCE", bold=True, size=12)

    for role in data["leadership_experience"]:
        role_para = doc.add_paragraph()
        set_paragraph_spacing(role_para, before=3.35, after=0)
        add_right_tab(role_para)
        add_run(role_para, role["organization"], bold=True, size=10)
        if role.get("role"):
            add_run(role_para, ", ", size=10)
            add_run(role_para, role["role"] + " ", italic=True, size=10)
        else:
            add_run(role_para, " ", size=10)
        add_run(role_para, "| " + role["location"], size=10)
        add_run(role_para, "\t", size=10)
        add_run(role_para, role["dates"], size=10)

        for bullet in role.get("bullets", []):
            bp = doc.add_paragraph()
            set_paragraph_spacing(bp, before=3.35, after=0)
            add_run(bp, bullet, size=10)

    # === Blank line ===
    blank4 = doc.add_paragraph()
    set_paragraph_spacing(blank4, before=0, after=0)

    # === SKILLS & INTERESTS ===
    skills_heading = doc.add_paragraph()
    set_paragraph_spacing(skills_heading, before=3.35, after=0)
    add_run(skills_heading, "SKILLS & INTERESTS", bold=True, size=12)

    si = data["skills_interests"]

    tech_para = doc.add_paragraph()
    set_paragraph_spacing(tech_para, before=3.35, after=0)
    add_run(tech_para, "Technical: ", italic=True, size=10)
    add_run(tech_para, si["technical"], size=10)

    lang_para = doc.add_paragraph()
    set_paragraph_spacing(lang_para, before=3.35, after=0)
    add_run(lang_para, "Languages: ", italic=True, size=10)
    add_run(lang_para, si["languages"], size=10)

    interests_para = doc.add_paragraph()
    set_paragraph_spacing(interests_para, before=3.35, after=0)
    add_run(interests_para, "Interests: ", italic=True, size=10)
    add_run(interests_para, si["interests"], size=10)

    # Save
    doc.save(output_path)
    print(f"Resume saved to {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} <input.json> <output.docx>")
        sys.exit(1)

    with open(sys.argv[1], 'r') as f:
        data = json.load(f)

    generate_resume(data, sys.argv[2])
